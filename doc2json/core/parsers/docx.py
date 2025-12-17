"""DOCX document parser for Microsoft Word files."""

import logging
import os

import docx

from doc2json.core.exceptions import ParserError

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parser for Microsoft Word (.docx) files.

    Extracts text from paragraphs and tables in Word documents.
    """

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self, include_tables: bool = True):
        """Initialize DOCX parser.

        Args:
            include_tables: Whether to extract text from tables (default True)
        """
        self.include_tables = include_tables

    def can_parse(self, file_path: str) -> bool:
        """Check if this is a DOCX file."""
        _, ext = os.path.splitext(file_path)
        return ext.lower() in self.SUPPORTED_EXTENSIONS

    def _extract_paragraphs(self, doc) -> list[str]:
        """Extract text from all paragraphs."""
        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)
        return texts

    def _extract_tables(self, doc) -> list[str]:
        """Extract text from all tables.

        Converts tables to a simple text representation with
        pipe-separated columns and newline-separated rows.
        """
        texts = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Remove duplicate cells (merged cells appear multiple times)
                seen = set()
                unique_cells = []
                for cell in cells:
                    if cell not in seen:
                        seen.add(cell)
                        unique_cells.append(cell)
                table_rows.append(" | ".join(unique_cells))

            if table_rows:
                texts.append("\n".join(table_rows))

        return texts

    def parse(self, file_path: str) -> str:
        """Parse a DOCX file and extract text.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content

        Raises:
            ParserError: If DOCX cannot be parsed
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            doc = docx.Document(file_path)

            # Extract paragraphs
            content_parts = self._extract_paragraphs(doc)

            # Extract tables if enabled
            if self.include_tables:
                table_texts = self._extract_tables(doc)
                if table_texts:
                    content_parts.extend(table_texts)

            return "\n\n".join(content_parts)

        except Exception as e:
            if "docx" in str(type(e).__module__):
                raise ParserError(f"Failed to parse DOCX: {e}")
            raise

    def get_metadata(self, file_path: str) -> dict:
        """Extract metadata from a DOCX file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = docx.Document(file_path)
        props = doc.core_properties

        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
            "last_modified_by": props.last_modified_by or "",
        }

    def analyze(self, file_path: str) -> dict:
        """Analyze a DOCX file structure."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = docx.Document(file_path)

        # Count paragraphs with content
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]

        # Count tables
        tables = doc.tables

        # Calculate total characters
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)

        return {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "total_characters": total_chars,
            "has_tables": len(tables) > 0,
        }
